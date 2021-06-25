from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import ListView
import sqlite3
import os
from .models import Generator, Category, Prepod, Discip, GodNabora, FormEducation, NapravPodgotovki, ZUV, ParsBook, \
    Competence, ParsComp, SelectComp, TargetsAndTasks, SelectBooks, Users, PrepInfo, PrepFIO, RecomendBoook
from .forms import NewsForm, TestForm, PrepForm, UserRegisterForm, UserLoginForm, ContactForm, BasicDataForm, BasicForm, \
    BDForm, GandO, PlanResEd, CompSelect, Books, UploadFileForm, StructDiscip, ROPHead, AddPrep
from django.contrib import messages
from django.contrib.auth import login, logout
from django.core.mail import send_mail, EmailMessage
import requests
from bs4 import BeautifulSoup
from collections import Counter
from openpyxl import load_workbook
import re
from docx import Document
from docxtpl import DocxTemplate
from rutermextract import TermExtractor


def register(request):
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            Users.objects.create(name=form.cleaned_data['username'], e_mail=form.cleaned_data['email'],
                                 role=form.cleaned_data['role_prep'])
            login(request, user)
            messages.success(request, 'Вы успешно зарегестрировались')
            return redirect('home')
        else:
            messages.error(request, 'Ошибка регистрации')
    else:
        form = UserRegisterForm()
    return render(request, 'generation_rpd/register.html', {"form": form})


def user_login(request):
    if request.method == 'POST':
        form = UserLoginForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('home')
    else:
        form = UserLoginForm()

    return render(request, 'generation_rpd/login.html', {"form": form})


def user_logout(request):
    logout(request)
    delete_sqlite_record()
    return redirect('login')


class HomeNews(ListView):
    model = Generator
    template_name = 'generation_rpd/home_news_list.html'
    context_object_name = 'news'

    # extra_context = {'title': 'Главная'}

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Главная страница'
        return context

    def get_queryset(self):
        return Generator.objects.filter(is_published=True)


class NewsByCategory(ListView):
    model = Generator
    template_name = 'generation_rpd/home_news_list.html'
    context_object_name = 'news'
    allow_empty = False

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = Category.objects.get(pk=self.kwargs['category_id'])
        return context

    def get_queryset(self):
        return Generator.objects.filter(category_id=self.kwargs['category_id'], is_published=True)


# def index(request):
#     news = Generator.objects.all()
#     categories = Category.objects.all()
#     context = {
#         'news': news,
#         'title': 'Список новостей',
#     }
#     return render(request, 'generation_rpd/index.html', context=context)


def test(request):
    if request.method == 'POST':
        email = EmailMessage(
            'Рабочая программа дисциплины',
            'Эта программа сделана с помощью конструктора рабочей программы дисциплины',
            'syava_test@mail.ru',
            ['sevostyanov1999@gmail.com'],
            ['syava_test@mail.ru']
        )

        email.attach_file('RPD.docx')
        email.send()
        return redirect('test')

    return render(request, 'generation_rpd/test.html')


def get_category(request, category_id):
    news = Generator.objects.filter(category_id=category_id)

    category = Category.objects.get(pk=category_id)
    return render(request, 'generation_rpd/category.html', {'news': news, 'category': category})


def view_news(request, news_id):
    # news_item = Generator.objects.get(pk=news_id)
    news_item = get_object_or_404(Generator, pk=news_id)
    return render(request, 'generation_rpd/view_news.html', {"news_item": news_item})


def add_news(request):
    if request.method == 'POST':
        form = NewsForm(request.POST)
        if form.is_valid():
            # print(form.cleaned_data)
            # news = Generator.objects.create(**form.cleaned_data)
            news = form.save()
            return redirect(news)
    else:
        form = NewsForm()
    return render(request, 'generation_rpd/add_news.html', {'form': form})


def add_prep(request):
    if request.method == 'POST':
        form = PrepForm(request.POST)
        if form.is_valid():
            # print(form.cleaned_data)
            # news = Generator.objects.create(**form.cleaned_data)
            prep = form.save()
            return redirect(prep)
    else:
        form = PrepForm()
    return render(request, 'generation_rpd/add_test.html', {'form': form})


def add_test(request):
    if request.method == 'POST':
        form = TestForm(request.POST)
        if form.is_valid():
            # print(form.cleaned_data)
            prep = Prepod.objects.create(**form.cleaned_data)
            return redirect('home')
    else:
        form = TestForm()
    return render(request, 'generation_rpd/add_test.html', {'form': form})


def add_basic_data(request):
    if request.method == 'POST':
        form = BasicForm(request.POST)
        if form.is_valid():
            # print(form.cleaned_data)
            discip = form.save()
            return redirect('home')
    else:
        form = BasicForm()
    return render(request, 'generation_rpd/add_basic_data.html', {'form': form})


def add_test_data(request):
    if request.method == 'POST':
        form = BDForm(request.POST)
        if form.is_valid():
            print(form.cleaned_data)
            delete_sqlite_record()
            #GodNabora.objects.create(god_nabora_num=form.cleaned_data['god_nabora_num'])
            #NapravPodgotovki.objects.create(naprav_podgotovki_title=form.cleaned_data['naprav_podgotovki_title'])
            #Discip.objects.create(discip_title=form.cleaned_data['discip_title'])
            #FormEducation.objects.create(form_education_title=form.cleaned_data['form_education_title'])
            return redirect('add_goals')
    else:
        form = BDForm()
    return render(request, 'generation_rpd/add_basic_data.html', {'form': form})


def add_goals(request):
    if request.method == 'POST':
        form = GandO(request.POST)
        if form.is_valid():
            TargetsAndTasks.objects.create(target=form.cleaned_data['g_a_o'], task=form.cleaned_data['task'],
                                           place_discip=form.cleaned_data['mesto_discip'])
            return redirect('add_plan_res_education')
    else:
        form = GandO()
    return render(request, 'generation_rpd/goals.html', {'form': form})


def opn_main_window(request):
    return render(request, 'generation_rpd/main_main.html')


def add_plan_res_ed(request):
    if request.method == 'POST':
        form = CompSelect(request.POST)
        if form.is_valid():
            print(form.cleaned_data)
            return redirect('home')
    else:
        form = CompSelect()

    bd1 = SelectComp.objects.all()
    return render(request, 'generation_rpd/add_compet.html', {'form': form, 'bd1': bd1})


def connect_db(request):
    conn = sqlite3.connect('db.sqlite3')
    cur = conn.cursor()
    cur.execute("SELECT * FROM generation_rpd_prepod")
    one_result = cur.fetchmany(10)
    print(one_result)
    return render(request, 'generation_rpd/add_plan_res_ed.html')


def bd_test(request):
    bd = Prepod.objects.all()
    return render(request, "generation_rpd/home_news_list.html", {'bd': bd})


def parsing_book(request):
    lis = list(range(1000, 1083))
    list_book = []
    list_book_ann = []
    for i in lis:
        link = f'https://www.iprbookshop.ru/10{i}.html'
        website_url = requests.get(link).text
        soup = BeautifulSoup(website_url, 'lxml')

        items = soup.find_all('div', {'class': 'col-sm-12'})
        res = items[2].get_text()
        items1 = soup.find('meta', property="og:description")

        ParsBook.objects.create(description_b=''.join(res.split('\r\n\t')), ann_b=items1['content'])

    return render(request, "generation_rpd/home_news_list.html")


def get_comp(request):
    """ Название файлов """
    filename1 = "comp.xlsx"

    """ Списки """
    spisok_zuv = []
    spisok_komp = []
    description_zuv = []

    """ Считываем данные из excel """
    wb = load_workbook(filename1)
    sheet = wb['Лист1']

    """ Загружаем ЗУВы и разбиваем их на слова, а также удаляем Знать, Уметь, Владеть """
    for row in sheet['C1':'C35']:
        zuv = ''
        for cell in row:
            zuv = zuv + str(cell.value)
            your_string = zuv
            removal_list = ['Знать', 'Уметь', 'Владеть', ' и ', 'для', 'None', ' в ', ' на ', ' по ', 'знать', 'Задача',
                            'Лабораторная', 'Практикум', 'Краткое', 'cодержание', 'Задачи', 'Вариант', 'Основные',
                            ' с ', ' к ', ' при ']
            for word in removal_list:
                your_string = your_string.replace(word, '')
        description_zuv.append(your_string)
        result_key = re.findall(r'\w+', your_string)
        col_count = Counter(result_key).most_common(5)
        result_main23 = re.sub(r'\d', '', str(col_count))
        res_res = re.findall(r'\w+', result_main23)
        spisok_zuv.append(res_res)  # Список слов из ЗУВ для каждой компетенции для сравнния с литературой
        Competence.objects.create(description_competence=your_string)

    for row in sheet['A1':'A35']:
        string = ''
        for cell in row:
            string = string + str(cell.value)
            your_string = string

        spisok_komp.append(your_string)
        Competence.objects.create(competence_text=your_string)
    return render(request, "generation_rpd/home_news_list.html")


def get_commp(request):
    doc = Document('09_04_03.docx')
    # par = doc.paragraphs
    tbl = doc.tables
    n = []
    listt = []
    l = []
    i = 0
    while i < 2:
        for strok in tbl[0].rows:
            rr = strok.cells[i].text.strip()
            listt.append(''.join(rr.split('\n')))
            # Competence.objects.create(description_competence=''.join(rr.split('\n')))
        l.append(listt)
        i += 1

    for i in listt:
        if i not in n:
            n.append(i)
            Competence.objects.create(description_competence=i)

    bd = Competence.objects.all()

    return render(request, "generation_rpd/add_plan_res_ed.html", {'bd': bd})


def pars_test(table):
    listt1 = []
    listt2 = []
    listt3 = []
    l = []
    i = 0
    for strok in table[0].rows:
        rr = strok.cells[0].text.strip()
        listt1.append(''.join(rr.split('\n')))
    listt1.pop(0)
    l.append(listt1)
    for strok in table[0].rows:
        rr = strok.cells[1].text.strip()
        listt2.append(''.join(rr.split('\n')))
    listt2.pop(0)
    l.append(listt2)
    for strok in table[0].rows:
        rr = strok.cells[2].text.strip()
        listt3.append(''.join(rr.split('\n')))
    listt3.pop(0)
    l.append(listt3)
    # for iitems in listt:
    #      if iitems not in ll1:
    #       ll1.append(iitems)
    return (l)


def get_data_comp(request):
    doc = Document('09_04_03.docx')
    # par = doc.paragraphs
    i = 0

    tbl = doc.tables

    list_ret = []
    k = pars_test(tbl)
    l1 = list(k[0][0::3])
    list_ret.append(l1)
    l2 = list(k[1][0::3])
    list_ret.append(l2)
    l1_1 = list(k[2][1::3])
    list_ret.append(l1_1)
    l1_2 = list(k[2][2::3])
    list_ret.append(l1_2)
    l1_3 = list(k[2][0::3])
    list_ret.append(l1_3)

    # while i < len(l1):
    #   ParsComp.objects.create(kod_comp=l1[i], descrip_comp=l2[i], kod_i_naim_comp1=l1_3[i], kod_i_naim_comp2=l1_2[i],
    #                          kod_i_naim_comp3=l1_1[i])
    # i += 1

    bd = ParsComp.objects.all()
    return render(request, "generation_rpd/plan_res.html", {'bd': bd})


def add_plan_res_education(request):
    if request.method == 'POST':
        form = CompSelect(request.POST)
        listt = []
        list_1 = []
        list_2 = []
        list_3 = []
        k = ParsComp.objects.values_list('descrip_comp')
        k1 = ParsComp.objects.values_list('kod_i_naim_comp1')
        k2 = ParsComp.objects.values_list('kod_i_naim_comp2')
        k3 = ParsComp.objects.values_list('kod_i_naim_comp3')

        for items in k:
            listt.append(list(items))

        for items in k1:
            list_1.append(list(items))

        for items in k2:
            list_2.append(list(items))

        for items in k3:
            list_3.append(list(items))

        if form.is_valid():
            flag = 0
            k = str(form.cleaned_data['comp'])
            # print(list_1)
            for itt in listt:
                for i in itt:
                    if i == k:
                        print(list_1[flag][0])
                        print(list_2[flag][0])
                        print(list_3[flag][0])
                        SelectComp.objects.create(descrip_c=k, kod_i_naim_c1=list_1[flag][0],
                                                  kod_i_naim_c2=list_3[flag][0],
                                                  kod_i_naim_c3=list_2[flag][0])

                        pass
                    flag += 1
            # if form.cleaned_data['comp'] ==
            return redirect('add_plan_res_education')
    else:
        form = CompSelect()

    bd1 = SelectComp.objects.all()
    print(bd1)
    return render(request, 'generation_rpd/add_compet.html', {'form': form, 'bd1': bd1})


def term_ex(term):
    term_extractor = TermExtractor()
    list_term = []
    for term in term_extractor(term, nested=True):
        list_term.append(term.normalized)
    set_term = set(list_term)
    return(set_term)


def term_comp():
    list_c = []
    res_l = []
    list_c_t = []

    bd1 = SelectComp.objects.values_list('descrip_c')
    bd11 = SelectComp.objects.values_list('kod_i_naim_c1')
    bd12 = SelectComp.objects.values_list('kod_i_naim_c2')
    bd13 = SelectComp.objects.values_list('kod_i_naim_c3')

    list_c.append(bd1)
    list_c.append(bd11)
    list_c.append(bd12)
    list_c.append(bd13)

    for items in list_c:
        for ii in items:
            res_l.append(*ii)

    for items_t_c in res_l:
        list_c_t.append(term_ex(items_t_c))

    return(list_c_t)




def term_book():
    list_b_t = []
    list_desc_book = []

    pb1 = ParsBook.objects.values_list('ann_b')

    for items_b in pb1:
        list_desc_book.append(*items_b)


    for items_t in list_desc_book:
        list_b_t.append(term_ex(items_t))

    return(list_b_t)


def con_db(cursor):
    conn = sqlite3.connect('/home/syava/webapp/db.sqlite3')
    cur = conn.cursor()
    cur.execute(cursor)
    one_result = cur.fetchall()
    list_res = []
    for itt in one_result:
        list_res.append(list(itt))
    # list_res[0][1]
    return(list_res)


def rec_book(list_1, list_2):
    b_a = []
    d_book = []
    ann = []
    desc = con_db('SELECT description_b FROM generation_rpd_parsbook')
    i = 0
    for items in list_1:
        res = list_2[3] & items
        if len(res) >= 1:
            b_a.append(res)
            d_book.append(desc[i][0])
            RecomendBoook.objects.create(desc=desc[i][0])
            # ann.append(list_res[i])
        i += 1
    return(d_book)


def struct_discip(request):
    if request.method == 'POST':
        form = StructDiscip(request.POST)
        if form.is_valid():
            t_comp = term_comp()
            t_book = term_book()

            print(rec_book(t_book, t_comp))
            return redirect('home')
    else:
        form = StructDiscip()

    bd1 = SelectComp.objects.all()
    return render(request, 'generation_rpd/struct_discip.html', {'form': form})


def delete_sqlite_record():
    try:
        sqlite_connection = sqlite3.connect('/home/syava/webapp/db.sqlite3')
        cursor = sqlite_connection.cursor()
        print("Подключен к SQLite")

        sql_update_query1 = "DELETE FROM generation_rpd_selectcomp"
        sql_update_query2 = "DELETE FROM generation_rpd_targetsandtasks"
        sql_update_query3 = "DELETE FROM generation_rpd_selectbooks"
        sql_update_query4 = "DELETE FROM generation_rpd_recomendboook"

        cursor.execute(sql_update_query1)
        cursor.execute(sql_update_query2)
        cursor.execute(sql_update_query3)
        cursor.execute(sql_update_query4)
        sqlite_connection.commit()

        print("Запись успешно удалена")
        cursor.close()

    except sqlite3.Error as error:
        print("Ошибка при работе с SQLite", error)
    finally:
        if sqlite_connection:
            sqlite_connection.close()
            print("Соединение с SQLite закрыто")


def gen_book(request):
    if request.method == 'POST':
        form = Books(request.POST)
        if form.is_valid():
            SelectBooks.objects.create(book=form.cleaned_data['an_book'])
            SelectBooks.objects.create(book=form.cleaned_data['offer_books'])
            print(form.cleaned_data)
            return redirect('gen_book')
    else:
        form = Books()
    bd_book = SelectBooks.objects.all()
    return render(request, 'generation_rpd/books.html', {'form': form, 'bd_book': bd_book})


def gen_final_doc():
    doc = DocxTemplate('tempp1.docx')
    # добавляем первый параграф

    list_1 = []
    list_11 = []
    list_12 = []
    list_13 = []
    list_2 = []
    list_3 = []

    bd1 = SelectComp.objects.values_list('descrip_c')
    bd11 = SelectComp.objects.values_list('kod_i_naim_c1')
    bd12 = SelectComp.objects.values_list('kod_i_naim_c2')
    bd13 = SelectComp.objects.values_list('kod_i_naim_c3')


    bd2 = TargetsAndTasks.objects.values_list('target')
    k = SelectBooks.objects.values_list('book')

    for items in k:
        list_1.append(items)

    for items in bd1:
        list_2.append(str(items))

    for items in bd11:
        list_11.append(str(items))

    for items in bd12:
        list_12.append(str(items))

    for items in bd13:
        list_13.append(str(items))

    for items in bd2:
        list_3.append(items)

    table_contents = []
    book_cont = []
    k = 0
    k1 = 0
    for i in list_2:
        table_contents.append({
            'cod': i,
            'desc1': list_11[k],
            'desc2': list_12[k],
            'desc3': list_13[k]
        })
        k += 1

    context = {
        'title': 'Automated Report',
        'table_contents': table_contents,
        'book': list_1
    }

    doc.render(context)
    # doc.add_paragraph(list_1[0])

    # добавляем еще два параграфа
    #par1 = doc.add_paragraph('Это второй абзац.')
    #par2 = doc.add_paragraph('Это третий абзац.')

    # добавляем текст во второй параграф
    # par1.add_run(list_2[0])
    #for items in k:
        #doc.add_paragraph(items)

    # добавляем текст в третий параграф
    # par2.add_run(list_3[0]).bold = True

    doc.save('RPD.docx')
    print(list_1[1])

def gen_res(request):
    bd_book = SelectBooks.objects.all()
    bd1 = SelectComp.objects.all()
    bd2 = TargetsAndTasks.objects.all()
    # print(bd_book)
    gen_final_doc()
    if request.method == 'POST':
        # print(bd_book)
        return redirect('gen_res')

    return render(request, 'generation_rpd/res.html', {'bd_book': bd_book, 'bd1': bd1, 'bd2': bd2})


def del_doc(request):
    bd_book = SelectBooks.objects.all()
    bd1 = SelectComp.objects.all()
    bd2 = TargetsAndTasks.objects.all()

    if request.method == 'POST':
        os.remove('/home/syava/webapp/RPD.docx')
        return redirect('add_test_data')
    return render(request, 'generation_rpd/res.html', {'bd_book': bd_book, 'bd1': bd1, 'bd2': bd2})


def get_discip():
    doc = Document('/home/syava/webapp/09_04_03.docx')
        # par = doc.paragraphs
    tbl = doc.tables

    listt = []
    l = []
    i = 0
    while i < 1:
        for strok in tbl[4].rows:
            rr = strok.cells[i].text.strip()
            if rr != '':
                listt.append(''.join(rr.split('\n')))
            #l.append(listt)
        i += 1
    for x in listt:
        while listt.count(x) > 1:
            listt.remove(x) # [1, 2, 3, 4]
    return(listt)


def rop_head(request):
    prep_info = PrepInfo.objects.all()
    if request.method == 'POST':
        form = ROPHead(request.POST)
        if form.is_valid():
            PrepInfo.objects.create(fio=form.cleaned_data['prep'], discip=form.cleaned_data['discip'])
            print(form.cleaned_data)
            # list_discip = get_discip()
            # for items in list_discip:
            #     Discip.objects.create(discip_title=items)

            return redirect('rop_head')

    else:
        form = ROPHead()

    return render(request, 'generation_rpd/rop_head.html', {'form': form, 'prep_info': prep_info})

def rop_add_fio(request):
    if request.method == 'POST':
        form = AddPrep(request.POST)
        if form.is_valid():
            PrepFIO.objects.create(fio=form.cleaned_data['fio'])
            print(form.cleaned_data)
            return redirect('rop_head')

    else:
        form = AddPrep()

    return render(request, 'generation_rpd/rop_add_fio.html', {'form': form})
